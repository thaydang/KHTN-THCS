"""Word document (.docx) export functionality for lesson plans.

This module converts lesson plan data to Word format using python-docx,
with support for embedding LaTeX formula images.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

from app.latex_renderer import LatexRenderer


class WordExporter:
    """Export lesson plans to Word (.docx) format."""

    def __init__(self, output_dir: Optional[Path] = None):
        """Initialize the Word exporter.

        Args:
            output_dir: Directory for temporary files (formula images, etc.)
        """
        self.output_dir = output_dir or Path("outputs")
        self.latex_renderer = LatexRenderer(output_dir=self.output_dir / "formulas")

    def export_lesson_plan(self, config: Dict[str, Any], output_path: Path) -> None:
        """Export a lesson plan configuration to a Word document.

        Args:
            config: Lesson plan configuration dictionary (same format as JSON input)
            output_path: Path where the Word document should be saved
        """
        doc = Document()

        # Set document properties
        self._set_document_properties(doc)

        # Add title and metadata
        self._add_metadata_section(doc, config.get("metadata", {}))

        # Add objectives
        if config.get("objectives"):
            self._add_bullet_section(doc, "Mục tiêu bài học", config["objectives"])

        # Add competencies
        if config.get("competencies"):
            self._add_bullet_section(
                doc, "Năng lực, phẩm chất hình thành", config["competencies"]
            )

        # Add materials
        if config.get("materials"):
            self._add_bullet_section(doc, "Học liệu và thiết bị", config["materials"])

        # Add digital resources
        if config.get("digital_resources"):
            self._add_bullet_section(
                doc, "Bài giảng điện tử và học liệu số", config["digital_resources"]
            )

        # Add formulas table
        if config.get("formulas"):
            self._add_formulas_table(doc, config["formulas"])

        # Add activities
        if config.get("activities"):
            self._add_activities_section(doc, config["activities"])

        # Add assessment
        if config.get("assessment"):
            self._add_bullet_section(doc, "Đánh giá", config["assessment"])

        # Add homework
        if config.get("homework"):
            self._add_bullet_section(
                doc, "Hướng dẫn học tập tiếp theo", config["homework"]
            )

        # Add reflection
        if config.get("reflection"):
            self._add_bullet_section(
                doc, "Ghi chú và tự đánh giá", config["reflection"]
            )

        # Save the document
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

    def _set_document_properties(self, doc: Document) -> None:
        """Set document-wide properties like font and spacing."""
        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(13)

        # Set Vietnamese font
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    def _add_metadata_section(self, doc: Document, metadata: Dict[str, Any]) -> None:
        """Add the title and metadata section."""
        # Title
        title = metadata.get("title", "Kế hoạch bài dạy Khoa học Tự nhiên")
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.runs[0].font.size = Pt(16)
        title_para.runs[0].font.bold = True

        # Metadata items
        metadata_items = [
            ("Ngày dạy", metadata.get("date")),
            ("Khối lớp", metadata.get("grade")),
            ("Chủ đề", metadata.get("unit")),
            ("Bài học", metadata.get("topic")),
            ("Giáo viên", metadata.get("teacher")),
            ("Trường", metadata.get("school")),
        ]

        for label, value in metadata_items:
            if value:
                para = doc.add_paragraph()
                para.add_run(f"{label}: ").bold = True
                para.add_run(str(value))

        # Add spacing
        doc.add_paragraph()

    def _add_heading(self, doc: Document, text: str, level: int = 1) -> None:
        """Add a heading with consistent formatting."""
        heading = doc.add_heading(text, level=level)
        heading.runs[0].font.name = "Times New Roman"
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    def _add_bullet_section(
        self, doc: Document, title: str, items: List[str]
    ) -> None:
        """Add a section with a heading and bullet points."""
        self._add_heading(doc, title, level=2)

        for item in items:
            if item:
                # Check for LaTeX expressions and replace with images
                item_with_images = self._process_latex_in_text(doc, item)
                if not item_with_images:  # No LaTeX found, add as text
                    doc.add_paragraph(item, style="List Bullet")

    def _add_formulas_table(self, doc: Document, formulas: List[Dict[str, Any]]) -> None:
        """Add a table showing formulas with LaTeX rendered as images."""
        self._add_heading(doc, "Công thức và ký hiệu sử dụng", level=2)

        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Ký hiệu"
        header_cells[1].text = "Diễn giải"
        header_cells[2].text = "Biểu thức"

        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Add formula rows
        for formula in formulas:
            row_cells = table.add_row().cells
            row_cells[0].text = formula.get("symbol", "")
            row_cells[1].text = formula.get("description", "")

            # Render LaTeX formula as image
            latex_expr = formula.get("latex", "")
            if latex_expr:
                try:
                    image_path = self.latex_renderer.render_to_file(latex_expr)
                    # Add image to cell
                    paragraph = row_cells[2].paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(str(image_path), width=Inches(2.0))
                except Exception as e:
                    # Fallback to text if rendering fails
                    row_cells[2].text = f"${latex_expr}$"

        doc.add_paragraph()  # Add spacing

    def _add_activities_section(
        self, doc: Document, activities: List[Dict[str, Any]]
    ) -> None:
        """Add the teaching activities section."""
        self._add_heading(doc, "Tiến trình dạy học", level=2)

        for activity in activities:
            # Activity title
            title = activity.get("title", "Hoạt động")
            self._add_heading(doc, title, level=3)

            # Duration
            if activity.get("duration"):
                para = doc.add_paragraph()
                para.add_run("Thời lượng: ").bold = True
                para.add_run(activity["duration"])

            # Goals
            if activity.get("goals"):
                para = doc.add_paragraph()
                para.add_run("Mục tiêu hoạt động:").bold = True
                for goal in activity["goals"]:
                    doc.add_paragraph(goal, style="List Bullet 2")

            # Steps
            if activity.get("steps"):
                para = doc.add_paragraph()
                para.add_run("Tiến trình:").bold = True
                for step in activity["steps"]:
                    actor = step.get("actor", "Giáo viên")
                    content = step.get("content", "")
                    para = doc.add_paragraph(style="List Bullet 2")
                    para.add_run(f"{actor}: ").bold = True
                    # Process LaTeX in content
                    self._add_text_with_latex(para, content)

            # Digital assets
            if activity.get("digital_assets"):
                para = doc.add_paragraph()
                para.add_run("Học liệu/Bài giảng điện tử:").bold = True
                for asset in activity["digital_assets"]:
                    doc.add_paragraph(asset, style="List Bullet 2")

            doc.add_paragraph()  # Add spacing

    def _process_latex_in_text(self, doc: Document, text: str) -> bool:
        """Process text containing LaTeX expressions and add as paragraph with images.

        Args:
            doc: Word document
            text: Text that may contain LaTeX expressions like $...$

        Returns:
            True if LaTeX was found and processed, False otherwise
        """
        # Find all LaTeX expressions in the text
        latex_pattern = r"\$([^\$]+)\$"
        matches = list(re.finditer(latex_pattern, text))

        if not matches:
            return False

        # Create a paragraph and add text with images
        para = doc.add_paragraph(style="List Bullet")
        last_end = 0

        for match in matches:
            # Add text before the LaTeX
            if match.start() > last_end:
                para.add_run(text[last_end : match.start()])

            # Render and add LaTeX image
            latex_expr = match.group(1)
            try:
                image_path = self.latex_renderer.render_to_file(latex_expr)
                run = para.add_run()
                run.add_picture(str(image_path), height=Inches(0.3))
            except Exception:
                # Fallback to text
                para.add_run(f"${latex_expr}$")

            last_end = match.end()

        # Add remaining text
        if last_end < len(text):
            para.add_run(text[last_end:])

        return True

    def _add_text_with_latex(self, paragraph, text: str) -> None:
        """Add text to an existing paragraph, replacing LaTeX with images.

        Args:
            paragraph: Word paragraph object
            text: Text that may contain LaTeX expressions
        """
        latex_pattern = r"\$([^\$]+)\$"
        matches = list(re.finditer(latex_pattern, text))

        if not matches:
            paragraph.add_run(text)
            return

        last_end = 0
        for match in matches:
            # Add text before LaTeX
            if match.start() > last_end:
                paragraph.add_run(text[last_end : match.start()])

            # Render and add LaTeX image
            latex_expr = match.group(1)
            try:
                image_path = self.latex_renderer.render_to_file(latex_expr)
                run = paragraph.add_run()
                run.add_picture(str(image_path), height=Inches(0.25))
            except Exception:
                paragraph.add_run(f"${latex_expr}$")

            last_end = match.end()

        # Add remaining text
        if last_end < len(text):
            paragraph.add_run(text[last_end:])


def export_to_word(config: Dict[str, Any], output_path: Path) -> None:
    """Convenience function to export a lesson plan configuration to Word.

    Args:
        config: Lesson plan configuration dictionary
        output_path: Path where the Word document should be saved
    """
    exporter = WordExporter()
    exporter.export_lesson_plan(config, output_path)
